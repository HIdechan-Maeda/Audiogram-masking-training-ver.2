#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MarkdownファイルをWord形式（.docx）に変換するスクリプト
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_japanese_font(run, font_name='Hiragino Sans'):
    """日本語フォントを設定（macOS用）"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)

def parse_markdown_to_docx(md_file_path, docx_file_path):
    """MarkdownファイルをWord形式に変換"""
    
    # Word文書を作成
    doc = Document()
    
    # デフォルトスタイルの設定（macOS用）
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Hiragino Sans'
    font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Hiragino Sans')
    style._element.rPr.rFonts.set(qn('w:ascii'), 'Hiragino Sans')
    style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Hiragino Sans')
    
    # 段落スタイルの設定
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    
    # Markdownファイルを読み込む
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 空行
        if not line:
            doc.add_paragraph()
            i += 1
            continue
        
        # 見出し1 (# タイトル)
        if line.startswith('# ') and not line.startswith('##'):
            title = line[2:].strip()
            p = doc.add_heading(title, level=1)
            for run in p.runs:
                set_japanese_font(run, 'Hiragino Sans')
                run.font.size = Pt(16)
                run.font.bold = True
            i += 1
            continue
        
        # 見出し2 (## セクション)
        if line.startswith('## ') and not line.startswith('###'):
            title = line[3:].strip()
            p = doc.add_heading(title, level=2)
            for run in p.runs:
                set_japanese_font(run, 'Hiragino Sans')
                run.font.size = Pt(14)
                run.font.bold = True
            i += 1
            continue
        
        # 見出し3 (### サブセクション)
        if line.startswith('### ') and not line.startswith('####'):
            title = line[4:].strip()
            p = doc.add_heading(title, level=3)
            for run in p.runs:
                set_japanese_font(run, 'Hiragino Sans')
                run.font.size = Pt(12)
                run.font.bold = True
            i += 1
            continue
        
        # 見出し4 (#### 小見出し)
        if line.startswith('#### '):
            title = line[5:].strip()
            p = doc.add_heading(title, level=4)
            for run in p.runs:
                set_japanese_font(run, 'Hiragino Sans')
                run.font.size = Pt(11)
                run.font.bold = True
            i += 1
            continue
        
        # 水平線 (---)
        if line.startswith('---'):
            p = doc.add_paragraph('─' * 50)
            for run in p.runs:
                set_japanese_font(run, 'Hiragino Sans')
            i += 1
            continue
        
        # テーブル（簡易版）
        if '|' in line and line.count('|') >= 2 and not line.startswith('|'):
            # テーブル行を収集
            table_lines = []
            j = i
            while j < len(lines) and '|' in lines[j]:
                table_lines.append(lines[j].rstrip())
                j += 1
            
            if len(table_lines) >= 2:
                # テーブルを作成
                headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                if len(headers) > 0 and headers[0]:
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    
                    # ヘッダー行
                    header_cells = table.rows[0].cells
                    for k, header in enumerate(headers):
                        header_cells[k].text = header
                        for paragraph in header_cells[k].paragraphs:
                            for run in paragraph.runs:
                                set_japanese_font(run, 'Hiragino Sans')
                                run.font.size = Pt(10)
                                run.font.bold = True
                    
                    # データ行
                    for row_line in table_lines[2:]:  # 区切り行をスキップ
                        if '|' in row_line and not row_line.startswith('|---'):
                            cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                            if len(cells) == len(headers):
                                row = table.add_row()
                                for k, cell_text in enumerate(cells):
                                    row.cells[k].text = cell_text
                                    for paragraph in row.cells[k].paragraphs:
                                        for run in paragraph.runs:
                                            set_japanese_font(run, 'Hiragino Sans')
                                            run.font.size = Pt(10)
            i = j
            continue
        
        # 通常の段落（太字、リストなどを処理）
        p = doc.add_paragraph()
        add_formatted_text(p, line)
        i += 1
    
    # Wordファイルを保存
    doc.save(docx_file_path)
    print(f"Wordファイルを作成しました: {docx_file_path}")

def add_formatted_text(paragraph, text):
    """フォーマットされたテキストを段落に追加"""
    if not text:
        return
    
    # リスト項目の処理
    if text.startswith('- '):
        text = text[2:]
        # リスト記号を追加
        run = paragraph.add_run('・')
        set_japanese_font(run, 'Hiragino Sans')
        run.font.size = Pt(10.5)
    
    # 太字 (**text**) の処理
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        
        if part.startswith('**') and part.endswith('**'):
            # 太字テキスト
            run = paragraph.add_run(part[2:-2])
            set_japanese_font(run, 'Hiragino Sans')
            run.font.size = Pt(10.5)
            run.font.bold = True
        else:
            # 通常テキスト
            run = paragraph.add_run(part)
            set_japanese_font(run, 'Hiragino Sans')
            run.font.size = Pt(10.5)

if __name__ == '__main__':
    md_file = 'HearSim_基礎聴覚評価学_シラバス_2年生.md'
    docx_file = 'HearSim_基礎聴覚評価学_シラバス_2年生.docx'
    
    try:
        parse_markdown_to_docx(md_file, docx_file)
        print("変換が完了しました。")
    except Exception as e:
        print(f"エラーが発生しました: {e}")
        import traceback
        traceback.print_exc()
